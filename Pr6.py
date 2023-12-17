import openpyxl
import docx
from docx.shared import Pt
wb = openpyxl.load_workbook("excel.xlsx")
ws = wb.active

list = []
for x in ws.values:
    list.append(x)

doc = docx.Document()
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(24)

i = 0
while i < len(list) - 1:
    doc.add_paragraph("Награждается " + list[i][0] + " " + list[i][1])
    doc.add_paragraph("Учащияся " + str(list[i][2]) + " класса образовательного учреждения " + list[i][3])
    doc.add_page_break()
    i += 1
doc.add_paragraph("Награждается " + list[i][0] + " " + list[i][1])
doc.add_paragraph("Учащияся " + str(list[i][2]) + " класса образовательного учреждения " + list[i][3])
doc.save("document.docx")